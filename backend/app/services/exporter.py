from __future__ import annotations

import csv
import io
import json
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _md_section(title: str, body: Optional[str]) -> str:
    body = (body or "").strip()
    if not body:
        return ""
    return f"## {title}\n\n{body}\n\n"


def manuscript_to_markdown(manuscript: Dict[str, Any]) -> str:
    parts = []
    if manuscript.get("title"):
        parts.append(f"# {manuscript.get('title')}\n\n")
    parts.append(_md_section("Abstract", manuscript.get("abstract")))
    parts.append(_md_section("Introduction", manuscript.get("introduction")))
    parts.append(_md_section("Methods", manuscript.get("methods")))
    parts.append(_md_section("Results", manuscript.get("results")))
    parts.append(_md_section("Evidence Comparison", manuscript.get("evidence_comparison")))
    parts.append(_md_section("Discussion", manuscript.get("discussion")))
    parts.append(_md_section("Conclusion", manuscript.get("conclusion")))

    refs = manuscript.get("references") or []
    if refs:
        parts.append("## References\n\n")
        for i, r in enumerate(refs, start=1):
            title = r.get("title") or "Untitled"
            year = r.get("year")
            journal = r.get("journal")
            doi = r.get("doi")
            pmid = r.get("pmid")
            meta = " ".join([x for x in [str(year) if year else None, journal] if x])
            tail = " | ".join([x for x in [f"DOI:{doi}" if doi else None, f"PMID:{pmid}" if pmid else None] if x])
            line = f"{i}. {title}"
            if meta:
                line += f" ({meta})"
            if tail:
                line += f" — {tail}"
            parts.append(line + "\n")
        parts.append("\n")

    return "".join(parts).strip() + "\n"


def references_to_json(manuscript: Dict[str, Any]) -> bytes:
    refs = manuscript.get("references") or []
    return (json.dumps(refs, indent=2, ensure_ascii=False) + "\n").encode("utf-8")


def references_to_csv(manuscript: Dict[str, Any]) -> bytes:
    refs = manuscript.get("references") or []
    out = io.StringIO()
    writer = csv.DictWriter(out, fieldnames=["title", "doi", "pmid", "year", "journal", "authors"])
    writer.writeheader()
    for r in refs:
        writer.writerow(
            {
                "title": r.get("title") or "",
                "doi": r.get("doi") or "",
                "pmid": r.get("pmid") or "",
                "year": r.get("year") or "",
                "journal": r.get("journal") or "",
                "authors": "; ".join(r.get("authors") or []),
            }
        )
    return out.getvalue().encode("utf-8")


def manuscript_to_docx(manuscript: Dict[str, Any]) -> bytes:
    doc = Document()
    if manuscript.get("title"):
        doc.add_heading(str(manuscript.get("title")), level=0)

    def add_section(h: str, body: Optional[str]):
        body = (body or "").strip()
        if not body:
            return
        doc.add_heading(h, level=1)
        for para in body.split("\n\n"):
            p = para.strip()
            if p:
                doc.add_paragraph(p)

    add_section("Abstract", manuscript.get("abstract"))
    add_section("Introduction", manuscript.get("introduction"))
    add_section("Methods", manuscript.get("methods"))
    add_section("Results", manuscript.get("results"))
    add_section("Evidence Comparison", manuscript.get("evidence_comparison"))
    add_section("Discussion", manuscript.get("discussion"))
    add_section("Conclusion", manuscript.get("conclusion"))

    refs = manuscript.get("references") or []
    if refs:
        doc.add_heading("References", level=1)
        for i, r in enumerate(refs, start=1):
            title = r.get("title") or "Untitled"
            doi = r.get("doi")
            pmid = r.get("pmid")
            year = r.get("year")
            journal = r.get("journal")
            line = f"{i}. {title}"
            meta = " ".join([x for x in [str(year) if year else None, journal] if x])
            if meta:
                line += f" ({meta})"
            tail = " | ".join([x for x in [f"DOI:{doi}" if doi else None, f"PMID:{pmid}" if pmid else None] if x])
            if tail:
                line += f" — {tail}"
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def manuscript_to_pdf(manuscript: Dict[str, Any]) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, title=str(manuscript.get("title") or "Manuscript"))
    styles = getSampleStyleSheet()

    story = []
    title = (manuscript.get("title") or "").strip()
    if title:
        story.append(Paragraph(title, styles["Title"]))
        story.append(Spacer(1, 12))

    def add_section(h: str, body: Optional[str]):
        body = (body or "").strip()
        if not body:
            return
        story.append(Paragraph(h, styles["Heading2"]))
        story.append(Spacer(1, 6))
        for para in body.split("\n\n"):
            p = para.strip()
            if p:
                story.append(Paragraph(p.replace("\n", "<br/>"), styles["BodyText"]))
                story.append(Spacer(1, 6))
        story.append(Spacer(1, 6))

    add_section("Abstract", manuscript.get("abstract"))
    add_section("Introduction", manuscript.get("introduction"))
    add_section("Methods", manuscript.get("methods"))
    add_section("Results", manuscript.get("results"))
    add_section("Evidence Comparison", manuscript.get("evidence_comparison"))
    add_section("Discussion", manuscript.get("discussion"))
    add_section("Conclusion", manuscript.get("conclusion"))

    refs = manuscript.get("references") or []
    if refs:
        story.append(Paragraph("References", styles["Heading2"]))
        story.append(Spacer(1, 6))
        for i, r in enumerate(refs, start=1):
            title = r.get("title") or "Untitled"
            year = r.get("year")
            journal = r.get("journal")
            doi = r.get("doi")
            pmid = r.get("pmid")
            meta = " ".join([x for x in [str(year) if year else None, journal] if x])
            tail = " | ".join([x for x in [f"DOI:{doi}" if doi else None, f"PMID:{pmid}" if pmid else None] if x])
            line = f"{i}. {title}"
            if meta:
                line += f" ({meta})"
            if tail:
                line += f" — {tail}"
            story.append(Paragraph(line, styles["BodyText"]))
            story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue()

